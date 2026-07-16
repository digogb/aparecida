"""
Tests for app/services/extractor.py and app/services/content_assembler.py
"""
from __future__ import annotations

import io

import pytest

from app.services.content_assembler import (
    SECTION_SEPARATOR,
    assemble,
    extract_body_section,
)
from app.services.extractor import (
    _consertar_numeros_rachados,
    _tabela_para_markdown,
    extract_docx,
    extract_file,
    extract_pdf,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_pdf(text: str) -> bytes:
    """Create a minimal PDF containing *text* using reportlab."""
    reportlab = pytest.importorskip("reportlab", reason="reportlab not installed")
    from reportlab.pdfgen import canvas as rl_canvas

    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf)
    c.drawString(72, 720, text)
    c.save()
    return buf.getvalue()


def _make_docx(text: str) -> bytes:
    """Create a minimal DOCX containing *text*."""
    from docx import Document

    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_table(cell_a: str, cell_b: str) -> bytes:
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = cell_a
    table.rows[0].cells[1].text = cell_b
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# extract_pdf
# ---------------------------------------------------------------------------

class TestExtractPdf:
    def test_returns_text_from_valid_pdf(self):
        pdf_bytes = _make_pdf("Conteudo juridico relevante")
        text, method = extract_pdf(pdf_bytes)
        assert "Conteudo juridico relevante" in text
        assert method == "pdfplumber"

    def test_method_is_pdfplumber_for_normal_pdf(self):
        pdf_bytes = _make_pdf("Texto de exemplo")
        _, method = extract_pdf(pdf_bytes)
        assert method == "pdfplumber"

    def test_returns_tuple_of_str_str(self):
        pdf_bytes = _make_pdf("Qualquer texto")
        result = extract_pdf(pdf_bytes)
        assert isinstance(result, tuple) and len(result) == 2
        assert isinstance(result[0], str)
        assert isinstance(result[1], str)

    def test_empty_pdf_returns_string(self):
        reportlab = pytest.importorskip("reportlab")
        from reportlab.pdfgen import canvas as rl_canvas

        buf = io.BytesIO()
        c = rl_canvas.Canvas(buf)
        c.save()
        text, method = extract_pdf(buf.getvalue())
        assert isinstance(text, str)

    def test_invalid_bytes_raises_or_returns_failed_gracefully(self):
        # extract_file (dispatcher) handles the exception; extract_pdf may raise
        _, _, status = extract_file("doc.pdf", b"not a pdf")
        assert status == "failed"


# ---------------------------------------------------------------------------
# Conserto de números monetários rachados (pdfplumber bug)
# ---------------------------------------------------------------------------

class TestConsertarNumerosRachados:
    """pdfplumber às vezes insere whitespace dentro de números pt-BR quando
    a célula é estreita. Sem conserto, a IA lê o valor errado e o cálculo
    do art. 125 quebra (caso real: 'R$ 9 39.166,94' lido como dois tokens)."""

    def test_digito_isolado_com_espaco(self):
        assert _consertar_numeros_rachados("R$ 9 39.166,94") == "R$ 939.166,94"

    def test_quebra_antes_do_ponto(self):
        assert _consertar_numeros_rachados("R$ 1 .146.164,56") == "R$ 1.146.164,56"

    def test_valor_ja_correto_permanece(self):
        assert _consertar_numeros_rachados("R$ 234.791,74") == "R$ 234.791,74"

    def test_multiplos_valores_na_mesma_string(self):
        entrada = "Total A: R$ 2 06.997,62 | Total B: R$ 939.166,94"
        esperado = "Total A: R$ 206.997,62 | Total B: R$ 939.166,94"
        assert _consertar_numeros_rachados(entrada) == esperado

    def test_celulas_vazias_preserva_traco(self):
        assert _consertar_numeros_rachados("R$ - R$ 4.965,01 R$ -") == "R$ - R$ 4.965,01 R$ -"

    def test_nao_afeta_texto_sem_moeda(self):
        texto = "art. 125 da Lei nº 14.133/2021"
        assert _consertar_numeros_rachados(texto) == texto


# ---------------------------------------------------------------------------
# Render de tabela como markdown (preserva relação coluna↔valor)
# ---------------------------------------------------------------------------

class TestTabelaParaMarkdown:

    def test_tabela_simples(self):
        tab = [["A", "B"], ["1", "2"], ["3", "4"]]
        md = _tabela_para_markdown(tab)
        # GFM: header | sep | corpo
        assert "| A | B |" in md
        assert "| --- | --- |" in md
        assert "| 1 | 2 |" in md
        assert "| 3 | 4 |" in md

    def test_celula_vazia_vira_traco(self):
        tab = [["A", "B"], ["1", None]]
        md = _tabela_para_markdown(tab)
        assert "| 1 | - |" in md

    def test_normaliza_pipe_em_celula(self):
        # '|' dentro da célula quebraria o markdown
        tab = [["X"], ["a|b"]]
        md = _tabela_para_markdown(tab)
        assert "a/b" in md
        assert "a|b" not in md.replace("| a", "")  # só os delimitadores têm |

    def test_aplica_conserto_de_numero_rachado_nas_celulas(self):
        # Caso real: célula vinda do pdfplumber com número quebrado
        tab = [["VALOR"], ["R$ 9 39.166,94"]]
        md = _tabela_para_markdown(tab)
        assert "R$ 939.166,94" in md
        assert "R$ 9 39.166,94" not in md

    def test_tabela_vazia(self):
        assert _tabela_para_markdown([]) == ""


# ---------------------------------------------------------------------------
# extract_docx
# ---------------------------------------------------------------------------

class TestExtractDocx:
    def test_extracts_paragraph_text(self):
        docx_bytes = _make_docx("Paragrafo principal do parecer")
        text, method = extract_docx(docx_bytes)
        assert "Paragrafo principal do parecer" in text
        assert method == "python_docx"

    def test_extracts_table_cells(self):
        docx_bytes = _make_docx_with_table("Coluna A", "Coluna B")
        text, _ = extract_docx(docx_bytes)
        assert "Coluna A" in text
        assert "Coluna B" in text

    def test_empty_docx_returns_empty_string(self):
        from docx import Document

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        text, method = extract_docx(buf.getvalue())
        assert text == ""
        assert method == "python_docx"

    def test_returns_tuple_str_str(self):
        docx_bytes = _make_docx("Qualquer")
        result = extract_docx(docx_bytes)
        assert isinstance(result, tuple) and len(result) == 2
        assert all(isinstance(v, str) for v in result)


# ---------------------------------------------------------------------------
# extract_file (dispatcher)
# ---------------------------------------------------------------------------

class TestExtractFile:
    def test_pdf_dispatch_returns_three_tuple(self):
        pdf_bytes = _make_pdf("PDF dispatch test")
        result = extract_file("documento.pdf", pdf_bytes)
        assert len(result) == 3

    def test_pdf_status_success_when_text_found(self):
        pdf_bytes = _make_pdf("Texto valido")
        text, method, status = extract_file("documento.pdf", pdf_bytes)
        assert status == "success"
        assert text

    def test_docx_dispatch(self):
        docx_bytes = _make_docx("Dispatch DOCX")
        text, method, status = extract_file("parecer.docx", docx_bytes)
        assert "Dispatch DOCX" in text
        assert method == "python_docx"
        assert status == "success"

    def test_txt_file(self):
        content = "Consulta juridica municipio"
        text, method, status = extract_file("consulta.txt", content.encode("utf-8"))
        assert "Consulta juridica municipio" in text
        assert status == "success"

    def test_txt_latin1_encoding(self):
        content = "Câmara Municipal"
        text, method, status = extract_file("nota.txt", content.encode("latin-1"))
        assert status == "success"
        assert isinstance(text, str)

    def test_unsupported_extension_returns_failed(self):
        text, method, status = extract_file("foto.png", b"\x89PNG\r\n\x1a\n")
        assert text == ""
        assert status == "failed"
        assert method is None

    def test_corrupted_pdf_returns_failed(self):
        text, method, status = extract_file("corrupted.pdf", b"not a pdf at all")
        assert status == "failed"
        assert text == ""

    def test_corrupted_docx_returns_failed(self):
        text, method, status = extract_file("bad.docx", b"not a docx")
        assert status == "failed"

    def test_status_values_are_valid(self):
        valid = {"success", "partial", "failed"}
        for name, data in [
            ("ok.docx", _make_docx("texto")),
            ("ok.txt", b"texto"),
            ("bad.png", b"\x89PNG"),
        ]:
            _, _, status = extract_file(name, data)
            assert status in valid


# ---------------------------------------------------------------------------
# content_assembler.assemble
# ---------------------------------------------------------------------------

class TestAssemble:
    def test_body_only(self):
        result = assemble("Corpo do email", [])
        assert "Corpo do email" in result

    def test_body_and_attachment(self):
        result = assemble("Corpo do email", ["Texto do anexo"])
        assert "Corpo do email" in result
        assert "Texto do anexo" in result
        assert "---" in result

    def test_multiple_attachments(self):
        result = assemble("Corpo", ["Anexo 1", "Anexo 2"])
        assert result.count("---") == 2

    def test_empty_body_with_attachment(self):
        result = assemble("", ["Apenas o anexo"])
        assert "Apenas o anexo" in result
        # O corpo (mesmo vazio) permanece como primeira seção: extract_body_section
        # separa corpo de anexos pelo primeiro separador.
        assert result.startswith(SECTION_SEPARATOR)
        assert extract_body_section(result) == ""

    def test_all_empty_returns_empty(self):
        result = assemble("", [])
        assert result == ""

    def test_skips_whitespace_only_attachments(self):
        result = assemble("Corpo", ["   \n\n  ", "Valido"])
        assert "Valido" in result
        # Only one separator between body and the valid attachment
        assert result.count("---") == 1

    def test_strips_email_headers(self):
        body = "De: alguem@example.com\nAssunto: Consulta\n\nConteudo real"
        result = assemble(body, [])
        assert "Conteudo real" in result
        assert "De:" not in result
        assert "Assunto:" not in result

    def test_strips_quoted_message_footer(self):
        body = "Minha pergunta\n\n--- Mensagem original ---\nTexto anterior"
        result = assemble(body, [])
        assert "Minha pergunta" in result
        assert "Mensagem original" not in result

    def test_sections_separated_by_rule(self):
        result = assemble("A", ["B"])
        assert result == "A\n\n---\n\nB"
