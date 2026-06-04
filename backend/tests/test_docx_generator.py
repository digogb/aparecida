"""
Tests for docx_generator (Camada 5).

Cobre:
- gerar_parecer_bytes: produz .docx válido, valida campos obrigatórios.
- IRR-1: ementa em CAPS com figure-dash.
- Marcadores [REVISAR — ...] e [!VERIFICAR: ... !]: vermelho/negrito inline.
- Blockquote (citação legal): left_indent=4cm, bold, sem prefixo '>' nem itálico markdown.
- Assinaturas com espaços manuais calibrados (Ione 25, Matheus/Flávio 13/26, Valéria 26).
- Parser minuta_from_tiptap: extrai dispositivo e alíneas.

Run: pytest backend/tests/test_docx_generator.py -v
"""
from __future__ import annotations

import io

import pytest
from docx import Document
from docx.shared import Cm, RGBColor

from app.services.docx_generator import (
    COR_MARCADOR_REVISAO,
    PADRAO_MARCADOR_QUALQUER,
    PADRAO_MARCADOR_REVISAR,
    PADRAO_MARCADOR_VERIFICAR,
    _ensure_relatorio_termina_com_passa_se,
    _identify_section,
    _normalizar_blockquote,
    _parse_conclusao,
    _parse_ementa_palavras_chave,
    contar_marcadores,
    gerar_parecer_bytes,
    minuta_from_tiptap,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

MINUTA_VALIDA = {
    "consulente": "Pregoeiro Oficial do Município de Potengi/CE — Pregão Eletrônico nº 2025.06.06.1",
    "ementa_palavras_chave": [
        "DIREITO ADMINISTRATIVO",
        "LEI Nº 14.133/2021",
        "PREGÃO ELETRÔNICO",
        "RECURSO ADMINISTRATIVO CONTRA INABILITAÇÃO",
        "PARECER PERTINENTE — PROVIMENTO RECOMENDADO",
    ],
    "relatorio_paragrafos": [
        "O Pregoeiro do Município de Potengi/CE submete à apreciação desta assessoria jurídica o Recurso Administrativo interposto pelo POPLAB.",
        "O POPLAB alega ter sido inabilitado indevidamente.",
        "É o breve relatório. Passa-se à fundamentação.",
    ],
    "fundamentos_paragrafos": [
        "O recurso administrativo em exame enquadra-se na hipótese do art. 165, I, da Lei nº 14.133/2021.",
        "No mérito, o ponto nuclear da controvérsia é fático.",
    ],
    "conclusao_dispositivo": "Diante do exposto, o parecer é no sentido de que o Recurso é PERTINENTE. Recomenda-se:",
    "recomendacoes_alineas": [
        ("a", "o conhecimento do recurso"),
        ("b", "no mérito, o provimento do recurso"),
    ],
    "data_extenso": "Fortaleza/CE, 15 de maio de 2026",
    "subtipo": "recurso",
    "vertente": "licitacao_14133",
}


def _runs_de_tudo(doc: Document):
    """Itera por todos os runs do documento (paragraphs principais)."""
    for para in doc.paragraphs:
        for run in para.runs:
            yield run


# ---------------------------------------------------------------------------
# gerar_parecer_bytes — saída e validações
# ---------------------------------------------------------------------------

class TestGerarParecerBytes:

    def test_retorna_bytes_zip_validos(self):
        bytes_docx = gerar_parecer_bytes(MINUTA_VALIDA)
        assert isinstance(bytes_docx, bytes)
        assert len(bytes_docx) > 0
        assert bytes_docx[:2] == b"PK"  # magic ZIP

    def test_docx_carregavel_por_python_docx(self):
        bytes_docx = gerar_parecer_bytes(MINUTA_VALIDA)
        doc = Document(io.BytesIO(bytes_docx))
        textos = [p.text for p in doc.paragraphs]
        assert "PARECER JURÍDICO" in textos
        assert any("I — RELATÓRIO" in t for t in textos)
        assert any("II — FUNDAMENTOS" in t for t in textos)
        assert any("III — CONCLUSÃO" in t for t in textos)

    def test_campo_obrigatorio_ausente_levanta_keyerror(self):
        minuta = dict(MINUTA_VALIDA)
        del minuta["fundamentos_paragrafos"]
        with pytest.raises(KeyError, match="fundamentos_paragrafos"):
            gerar_parecer_bytes(minuta)

# ---------------------------------------------------------------------------
# IRR-1 — Ementa em CAPS com figure-dash
# ---------------------------------------------------------------------------

class TestEmentaCAPS:

    def test_ementa_em_maiusculas(self):
        bytes_docx = gerar_parecer_bytes(MINUTA_VALIDA)
        doc = Document(io.BytesIO(bytes_docx))
        ementa_para = next(p for p in doc.paragraphs if p.text.startswith("EMENTA"))
        # Texto inteiro está em CAPS (exceto sinais e dígitos)
        assert ementa_para.text == ementa_para.text.upper()

    def test_ementa_usa_figure_dash(self):
        bytes_docx = gerar_parecer_bytes(MINUTA_VALIDA)
        doc = Document(io.BytesIO(bytes_docx))
        ementa_para = next(p for p in doc.paragraphs if p.text.startswith("EMENTA"))
        # Figure-dash U+2015 é o separador
        assert "―" in ementa_para.text or "―" in ementa_para.text

    def test_ementa_termina_com_ponto(self):
        bytes_docx = gerar_parecer_bytes(MINUTA_VALIDA)
        doc = Document(io.BytesIO(bytes_docx))
        ementa_para = next(p for p in doc.paragraphs if p.text.startswith("EMENTA"))
        assert ementa_para.text.endswith(".")


# ---------------------------------------------------------------------------
# Marcadores de revisão humana
# ---------------------------------------------------------------------------

class TestMarcadoresRevisar:

    def test_padrao_revisar_captura_basico(self):
        texto = "Conforme [REVISAR — confirmar art. 53 da Lei nº 14.133/2021] o gestor."
        m = PADRAO_MARCADOR_REVISAR.search(texto)
        assert m is not None
        assert "art. 53" in m.group(0)

    def test_padrao_revisar_tolera_hifen_e_travessao(self):
        for sep in ("-", "–", "—"):
            texto = f"[REVISAR {sep} VERIFICAR]"
            assert PADRAO_MARCADOR_REVISAR.search(texto) is not None

    def test_marcador_revisar_em_paragrafo_renderiza_vermelho(self):
        minuta = dict(MINUTA_VALIDA)
        minuta["fundamentos_paragrafos"] = [
            "O dispositivo aplicável é [REVISAR — art. exato] da lei.",
        ]
        bytes_docx = gerar_parecer_bytes(minuta)
        doc = Document(io.BytesIO(bytes_docx))

        vermelhos = []
        for run in _runs_de_tudo(doc):
            cor = run.font.color.rgb
            if cor is not None and cor == COR_MARCADOR_REVISAO:
                vermelhos.append(run.text)

        assert any("REVISAR" in t for t in vermelhos), (
            f"Esperado run vermelho com 'REVISAR'. Vermelhos: {vermelhos!r}"
        )


class TestMarcadoresVerificar:

    def test_padrao_verificar_captura(self):
        texto = "Quanto à norma [!VERIFICAR: Lei Municipal 1.234/2018 !] não localizada."
        m = PADRAO_MARCADOR_VERIFICAR.search(texto)
        assert m is not None

    def test_marcador_verificar_em_paragrafo_renderiza_vermelho(self):
        minuta = dict(MINUTA_VALIDA)
        minuta["fundamentos_paragrafos"] = [
            "A norma de regência é [!VERIFICAR: art. X da LC municipal !] desta cidade.",
        ]
        bytes_docx = gerar_parecer_bytes(minuta)
        doc = Document(io.BytesIO(bytes_docx))

        vermelhos = [
            run.text for run in _runs_de_tudo(doc)
            if run.font.color.rgb is not None
            and run.font.color.rgb == COR_MARCADOR_REVISAO
        ]
        assert any("VERIFICAR" in t for t in vermelhos)


class TestBlockquote:

    def test_normalizar_remove_prefixo_e_italico(self):
        # `>` (markdown blockquote) e `*..*` (itálico) são strippados; aspas
        # eventualmente presentes na citação são preservadas (decisão editorial).
        bruto = '> *"Art. 168. O recurso terá efeito suspensivo."*'
        assert _normalizar_blockquote(bruto) == '"Art. 168. O recurso terá efeito suspensivo."'

    def test_normalizar_multilinha(self):
        bruto = "> linha 1\n> linha 2\n> linha 3"
        assert _normalizar_blockquote(bruto) == "linha 1 linha 2 linha 3"

    def test_normalizar_sem_italico(self):
        assert _normalizar_blockquote("> texto direto") == "texto direto"

    def test_blockquote_no_render_tem_left_indent_4cm(self):
        minuta = dict(MINUTA_VALIDA)
        minuta["fundamentos_paragrafos"] = [
            "Parágrafo normal de fundamentos.",
            "> Art. 168. O recurso terá efeito suspensivo.",
            "Após a citação, retoma-se o raciocínio.",
        ]
        bytes_docx = gerar_parecer_bytes(minuta)
        doc = Document(io.BytesIO(bytes_docx))

        bq = next(p for p in doc.paragraphs if "Art. 168" in p.text)
        # Sem prefixo '>' literal no .docx final
        assert not bq.text.startswith(">")
        # left_indent=4cm (parágrafo inteiro recuado, não só 1ª linha).
        # Tolerância para arredondamento twips↔EMU do python-docx.
        assert abs(bq.paragraph_format.left_indent.cm - 4.0) < 0.01
        assert bq.paragraph_format.first_line_indent in (None, Cm(0))
        # Bold
        assert all(r.font.bold for r in bq.runs if r.text.strip())


class TestContarMarcadores:

    def test_zero_marcadores(self):
        assert contar_marcadores("Texto sem marcador algum.") == 0
        assert contar_marcadores("") == 0
        assert contar_marcadores(None) == 0  # type: ignore

    def test_um_marcador(self):
        assert contar_marcadores("Foo [REVISAR — bar] baz.") == 1

    def test_mistura_dos_dois_tipos(self):
        texto = "[REVISAR — A] meio [!VERIFICAR: B !] fim [REVISAR — C]."
        assert contar_marcadores(texto) == 3

    def test_padrao_qualquer_captura_ambos(self):
        texto = "[REVISAR — A] e [!VERIFICAR: B !]"
        matches = PADRAO_MARCADOR_QUALQUER.findall(texto)
        assert len(matches) == 2


# ---------------------------------------------------------------------------
# Bloco de assinaturas — espaços calibrados
# ---------------------------------------------------------------------------

class TestAssinaturas:

    def test_ione_tem_25_espacos_avanco(self):
        bytes_docx = gerar_parecer_bytes(MINUTA_VALIDA)
        doc = Document(io.BytesIO(bytes_docx))
        # Procura o parágrafo de nome de Ione (primeiro run = 25 espaços, segundo = nome)
        para_ione = next(
            p for p in doc.paragraphs
            if p.text.strip() == "FRANCISCO IONE PEREIRA LIMA"
        )
        assert para_ione.runs[0].text == " " * 25
        assert para_ione.runs[1].text == "FRANCISCO IONE PEREIRA LIMA"

    def test_valeria_tem_26_espacos_avanco(self):
        bytes_docx = gerar_parecer_bytes(MINUTA_VALIDA)
        doc = Document(io.BytesIO(bytes_docx))
        para_val = next(
            p for p in doc.paragraphs
            if p.text.strip() == "VALÉRIA MATIAS DE ALENCAR"
        )
        assert para_val.runs[0].text == " " * 26
        assert para_val.runs[1].text == "VALÉRIA MATIAS DE ALENCAR"

    def test_matheus_flavio_em_paragrafo_unico_com_13_espacos(self):
        bytes_docx = gerar_parecer_bytes(MINUTA_VALIDA)
        doc = Document(io.BytesIO(bytes_docx))
        para_mf = next(
            p for p in doc.paragraphs
            if "MATHEUS" in p.text and "FLÁVIO" in p.text
        )
        # Três runs: nome_M + espaços + nome_F
        textos_runs = [r.text for r in para_mf.runs]
        assert textos_runs[0] == "MATHEUS NOGUEIRA PEREIRA LIMA"
        assert textos_runs[1] == " " * 13
        assert textos_runs[2] == "FLÁVIO HENRIQUE LUNA SILVA"

    def test_oabs_matheus_flavio_com_26_espacos(self):
        bytes_docx = gerar_parecer_bytes(MINUTA_VALIDA)
        doc = Document(io.BytesIO(bytes_docx))
        para_oabs = next(
            p for p in doc.paragraphs
            if "OAB-CE nº 31.251" in p.text and "OAB-CE nº 31.252" in p.text
        )
        textos = [r.text for r in para_oabs.runs]
        assert textos[0] == "OAB-CE nº 31.251"
        assert textos[1] == " " * 26
        assert textos[2] == "OAB-CE nº 31.252"


# ---------------------------------------------------------------------------
# Parser TipTap → minuta
# ---------------------------------------------------------------------------

class TestParseEmenta:

    def test_strip_prefix_e_ponto(self):
        textos = ["EMENTA: A ― B ― C."]
        palavras = _parse_ementa_palavras_chave(textos)
        assert palavras == ["A", "B", "C"]

    def test_sem_prefix(self):
        textos = ["A ― B ― C"]
        palavras = _parse_ementa_palavras_chave(textos)
        assert palavras == ["A", "B", "C"]

    def test_aceita_em_dash(self):
        textos = ["A — B — C"]
        palavras = _parse_ementa_palavras_chave(textos)
        assert palavras == ["A", "B", "C"]

    def test_vazio(self):
        assert _parse_ementa_palavras_chave([]) == []
        assert _parse_ementa_palavras_chave([""]) == []


class TestParseConclusao:

    def test_dispositivo_e_alineas(self):
        textos = [
            "Diante do exposto, o parecer é favorável. Recomenda-se:",
            "(a) primeira recomendação",
            "(b) segunda recomendação",
        ]
        dispositivo, alineas = _parse_conclusao(textos)
        # dispositivo agora é bloco {'text','attrs'} (carrega recuos da régua)
        assert "Diante do exposto" in dispositivo["text"]
        assert alineas == [
            ("a", "primeira recomendação"),
            ("b", "segunda recomendação"),
        ]

    def test_alineas_sem_parenteses(self):
        textos = ["Dispositivo.", "a) recomendação A", "b) recomendação B"]
        _, alineas = _parse_conclusao(textos)
        assert alineas == [("a", "recomendação A"), ("b", "recomendação B")]

    def test_paragrafo_pos_alineas_e_descartado(self):
        # Advertências terminais ao gestor foram abolidas; mesmo que escapem
        # do prompt, o parser as descarta silenciosamente.
        textos = [
            "Diante do exposto, parecer favorável.",
            "(a) recomendação",
            "Cumpre advertir o gestor sobre o risco de responsabilização.",
        ]
        dispositivo, alineas = _parse_conclusao(textos)
        assert "Diante do exposto" in dispositivo["text"]
        assert "Cumpre advertir" not in dispositivo["text"]
        assert alineas == [("a", "recomendação")]


class TestEnsureRelatorioFormula:

    def test_adiciona_quando_falta(self):
        paragrafos = ["Parágrafo 1.", "Parágrafo 2."]
        result = _ensure_relatorio_termina_com_passa_se(paragrafos)
        # a fórmula anexada vem como bloco {'text','attrs'} (padrão da casa)
        assert result[-1]["text"] == "É o breve relatório. Passa-se à fundamentação."

    def test_nao_duplica(self):
        paragrafos = ["Parágrafo 1.", "É o breve relatório. Passa-se à fundamentação."]
        result = _ensure_relatorio_termina_com_passa_se(paragrafos)
        assert len(result) == 2

    def test_vazio_permanece_vazio(self):
        assert _ensure_relatorio_termina_com_passa_se([]) == []


class TestIdentifySection:

    def test_variantes_de_hifen(self):
        for titulo in ("I — RELATÓRIO", "I – RELATÓRIO", "I - RELATÓRIO"):
            assert _identify_section(titulo) == "relatorio"

    def test_nome_simples_sem_numeral(self):
        assert _identify_section("FUNDAMENTOS") == "fundamentos"
        assert _identify_section("CONCLUSÃO") == "conclusao"

    def test_nao_reconhecido(self):
        assert _identify_section("INTRODUÇÃO") is None


# ---------------------------------------------------------------------------
# minuta_from_tiptap — caminho completo
# ---------------------------------------------------------------------------

def _tiptap_canonico() -> dict:
    """TipTap completo no padrão produzido pelo parecer_engine."""
    return {
        "type": "doc",
        "content": [
            {"type": "heading", "attrs": {"level": 2},
             "content": [{"type": "text", "text": "EMENTA"}]},
            {"type": "paragraph",
             "content": [{"type": "text", "text": "DIREITO ADMINISTRATIVO ― LEI 14.133/2021 ― PARECER FAVORÁVEL."}]},

            {"type": "heading", "attrs": {"level": 2},
             "content": [{"type": "text", "text": "I — RELATÓRIO"}]},
            {"type": "paragraph",
             "content": [{"type": "text", "text": "O órgão consulente submete consulta sobre licitação."}]},
            {"type": "paragraph",
             "content": [{"type": "text", "text": "É o breve relatório. Passa-se à fundamentação."}]},

            {"type": "heading", "attrs": {"level": 2},
             "content": [{"type": "text", "text": "II — FUNDAMENTOS"}]},
            {"type": "paragraph",
             "content": [{"type": "text", "text": "A análise da hipótese converge para a regularidade da contratação."}]},
            {"type": "paragraph",
             "content": [{"type": "text", "text": "Os requisitos legais foram atendidos."}]},

            {"type": "heading", "attrs": {"level": 2},
             "content": [{"type": "text", "text": "III — CONCLUSÃO"}]},
            {"type": "paragraph",
             "content": [{"type": "text", "text": "Diante do exposto, o parecer é FAVORÁVEL. Recomenda-se:"}]},
            {"type": "paragraph",
             "content": [{"type": "text", "text": "a) a observância do prazo legal"}]},
            {"type": "paragraph",
             "content": [{"type": "text", "text": "b) a publicação do extrato no diário oficial"}]},
        ],
    }


class TestMinutaFromTipTap:

    def test_minuta_completa(self):
        minuta = minuta_from_tiptap(
            _tiptap_canonico(),
            consulente="Pregoeiro de Potengi/CE",
            data_extenso="Fortaleza/CE, 15 de maio de 2026",
            subtipo="parecer",
            vertente="licitacao_14133",
        )
        assert minuta["consulente"] == "Pregoeiro de Potengi/CE"
        assert minuta["ementa_palavras_chave"] == [
            "DIREITO ADMINISTRATIVO",
            "LEI 14.133/2021",
            "PARECER FAVORÁVEL",
        ]
        assert len(minuta["relatorio_paragrafos"]) == 2
        # minuta_from_tiptap devolve blocos {'text','attrs'} (recuos por parágrafo da régua)
        assert minuta["relatorio_paragrafos"][-1]["text"].startswith("É o breve relatório")
        assert len(minuta["fundamentos_paragrafos"]) == 2
        assert minuta["conclusao_dispositivo"]["text"].startswith("Diante do exposto")
        assert minuta["recomendacoes_alineas"] == [
            ("a", "a observância do prazo legal"),
            ("b", "a publicação do extrato no diário oficial"),
        ]
        assert "advertencia_protetiva" not in minuta

    def test_tiptap_vazio_devolve_estrutura_minima(self):
        minuta = minuta_from_tiptap(
            {},
            consulente="X",
            data_extenso="Fortaleza/CE, 1 de janeiro de 2026",
        )
        assert minuta["ementa_palavras_chave"] == []
        assert minuta["relatorio_paragrafos"] == []
        assert minuta["fundamentos_paragrafos"] == []
        assert minuta["recomendacoes_alineas"] == []
        # consulente + data_extenso preservados
        assert minuta["consulente"] == "X"

    def test_minuta_alimenta_gerador_sem_erro(self):
        """End-to-end: TipTap canônico → minuta → bytes ZIP válido."""
        minuta = minuta_from_tiptap(
            _tiptap_canonico(),
            consulente="Pregoeiro de Potengi/CE",
            data_extenso="Fortaleza/CE, 15 de maio de 2026",
        )
        bytes_docx = gerar_parecer_bytes(minuta)
        assert bytes_docx[:2] == b"PK"
        doc = Document(io.BytesIO(bytes_docx))
        textos = "\n".join(p.text for p in doc.paragraphs)
        assert "DIREITO ADMINISTRATIVO" in textos
        assert "É o parecer, submetido à superior consideração." in textos


# ---------------------------------------------------------------------------
# Recuos por parágrafo (régua funcional do editor)
# ---------------------------------------------------------------------------

class TestRecuosPorParagrafo:
    """Os recuos definidos na régua do editor (atributos firstLineIndent/leftIndent/
    rightIndent no node paragraph) fluem até o paragraph_format do DOCX. Parágrafos
    sem atributos seguem o gabarito calibrado da casa (1ª linha 4cm, esq/dir não setados)."""

    def _doc_de_tiptap(self, paragraph_attrs: dict) -> Document:
        tiptap = {
            "type": "doc",
            "content": [
                {"type": "heading", "attrs": {"level": 2},
                 "content": [{"type": "text", "text": "II — FUNDAMENTOS"}]},
                {"type": "paragraph", "attrs": paragraph_attrs,
                 "content": [{"type": "text", "text": "Parágrafo com recuo custom da régua."}]},
                {"type": "paragraph",
                 "content": [{"type": "text", "text": "Parágrafo sem recuo segue o padrão."}]},
            ],
        }
        minuta = minuta_from_tiptap(
            tiptap, consulente="X", data_extenso="Fortaleza/CE, 1 de janeiro de 2026"
        )
        return Document(io.BytesIO(gerar_parecer_bytes(minuta)))

    def _para(self, doc: Document, trecho: str):
        for p in doc.paragraphs:
            if trecho in p.text:
                return p
        raise AssertionError(f"parágrafo não encontrado: {trecho!r}")

    def test_recuos_custom_refletem_no_docx(self):
        doc = self._doc_de_tiptap(
            {"firstLineIndent": 2, "leftIndent": 1.5, "rightIndent": 0.5}
        )
        pf = self._para(doc, "recuo custom da régua").paragraph_format
        # round-trip EMU→twips no XML introduz ~±0,0003cm; comparar com tolerância
        assert pf.first_line_indent.cm == pytest.approx(2, abs=0.01)
        assert pf.left_indent.cm == pytest.approx(1.5, abs=0.01)
        assert pf.right_indent.cm == pytest.approx(0.5, abs=0.01)

    def test_paragrafo_sem_attrs_mantem_padrao_casa(self):
        doc = self._doc_de_tiptap(
            {"firstLineIndent": 2, "leftIndent": 1.5, "rightIndent": 0.5}
        )
        pf = self._para(doc, "sem recuo segue o padrão").paragraph_format
        assert pf.first_line_indent.cm == pytest.approx(4, abs=0.01)  # padrão da casa
        assert pf.left_indent is None
        assert pf.right_indent is None

    def test_first_line_indent_zero_explicito_e_respeitado(self):
        # 0 explícito (usuário arrastou o marcador até 0) ≠ None (que herda 4cm)
        doc = self._doc_de_tiptap({"firstLineIndent": 0})
        pf = self._para(doc, "recuo custom da régua").paragraph_format
        assert pf.first_line_indent.cm == pytest.approx(0, abs=0.01)
