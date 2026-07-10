"""
Text extraction from PDF, DOCX, and plain-text files.

extract_pdf  -> (text, method)  — pdfplumber primary, Tesseract OCR fallback
extract_docx -> (text, method)  — python-docx paragraphs + tables
extract_file -> (text, method, status) — dispatcher by file extension
"""
from __future__ import annotations

import io
import logging
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Optional, Tuple

import chardet
import filetype
import pdfplumber
import pytesseract
from docx import Document

logger = logging.getLogger(__name__)

# LibreOffice conversion timeout (soffice can hang on malformed input)
_SOFFICE_TIMEOUT_S = 60


def _sanitize_text(text: str) -> str:
    """Strip NUL bytes — PostgreSQL TEXT columns reject 0x00 and abort the INSERT."""
    if not text:
        return text
    return text.replace("\x00", "")


def _extract_doc_via_libreoffice(file_bytes: bytes) -> str:
    """Convert legacy .doc (binary OLE) to text using `soffice --headless`.

    Uses an isolated user profile so concurrent calls don't contend on the
    default LibreOffice lock. Raises RuntimeError on any conversion failure.
    """
    with tempfile.TemporaryDirectory(prefix="ione_doc_") as workdir:
        input_path = Path(workdir) / "input.doc"
        input_path.write_bytes(file_bytes)

        try:
            subprocess.run(
                [
                    "soffice",
                    f"-env:UserInstallation=file://{workdir}/profile",
                    "--headless",
                    "--convert-to", "txt:Text",
                    "--outdir", workdir,
                    str(input_path),
                ],
                check=True,
                capture_output=True,
                timeout=_SOFFICE_TIMEOUT_S,
            )
        except subprocess.TimeoutExpired:
            raise RuntimeError(f"LibreOffice timed out after {_SOFFICE_TIMEOUT_S}s")
        except subprocess.CalledProcessError as exc:
            stderr = exc.stderr.decode("utf-8", errors="replace")[:300]
            raise RuntimeError(f"LibreOffice failed (rc={exc.returncode}): {stderr}")

        output_path = Path(workdir) / "input.txt"
        if not output_path.exists():
            raise RuntimeError("LibreOffice produced no .txt output")

        return output_path.read_text(encoding="utf-8", errors="replace")


# Número monetário em pt-BR após "R$": dígitos e pontos possivelmente entrecortados
# por espaços, seguidos de vírgula e 2 dígitos. Captura o miolo para limpar o
# whitespace que pdfplumber insere quando rachado entre células estreitas.
_RE_MONEY_BR = re.compile(r"R\$\s*([\d\.\s]+,\d{2})")


def _consertar_numeros_rachados(texto: str) -> str:
    """Une dígitos rachados por whitespace dentro de números monetários
    formatados em pt-BR (`R$ 9 39.166,94` → `R$ 939.166,94`,
    `R$ 1 .146.164,56` → `R$ 1.146.164,56`).

    pdfplumber às vezes insere espaço entre dígitos quando uma célula é
    estreita ou rendered em duas linhas. Sem isso o LLM lê o número
    errado e a conta do art. 125 vira fanfarra.
    """
    def _strip_spaces(m):
        miolo = re.sub(r"\s+", "", m.group(1))
        return f"R$ {miolo}"
    return _RE_MONEY_BR.sub(_strip_spaces, texto)


# Linha totalizadora do rodapé das planilhas de aditivo. É a fonte canônica do
# valor do contrato (e dos acréscimos/supressões), mas costuma ficar na última
# página — longe do início e, portanto, vulnerável a truncamento por limite de
# tokens. Detectamos e içamos ao topo do texto.
_RE_TOTALIZADOR = re.compile(r"VALOR\s+TOTAL\s+GERAL", re.IGNORECASE)


def resumo_financeiro_canonico(texto: str) -> str:
    """Extrai a(s) linha(s) totalizadora(s) ('VALOR TOTAL GERAL ...') e devolve
    um bloco rotulado para ser colocado no TOPO do texto.

    Retorna "" quando não há totalizador. O bloco resultante começa com o marcador
    `[RESUMO FINANCEIRO` para poder ser reidentificado e reordenado adiante no
    pipeline (antes de truncamentos por limite de tokens).
    """
    linhas: list[str] = []
    for raw in texto.splitlines():
        linha = raw.strip()
        if linha and _RE_TOTALIZADOR.search(linha) and linha not in linhas:
            linhas.append(linha)
    if not linhas:
        return ""
    return (
        "[RESUMO FINANCEIRO — totalizador do rodapé da planilha. "
        "Estes são os valores CANÔNICOS do contrato; preferir a qualquer valor "
        "solto no cabeçalho das páginas.]\n" + "\n".join(linhas)
    )


def _separar_cabecalho_repetido(
    page_texts: list[str], *, min_fraction: float = 0.4, min_pages: int = 3
) -> Tuple[list[str], list[str]]:
    """Remove linhas de cabeçalho/rodapé que se repetem em várias páginas.

    Em planilhas multipágina o mesmo cabeçalho-resumo aparece em toda página,
    inflando o texto e fazendo valores soltos (ex.: o teto de 25%) ocorrerem
    dezenas de vezes — o que enviesa o LLM. Linhas presentes em pelo menos
    `min_fraction` das páginas (e ao menos `min_pages`) são consideradas
    boilerplate: removidas do corpo e devolvidas UMA vez à parte.

    Returns (page_texts_limpos, linhas_boilerplate_em_ordem).
    """
    from collections import Counter

    n = len(page_texts)
    if n < min_pages:
        return page_texts, []

    threshold = max(min_pages, int(n * min_fraction))
    freq: Counter[str] = Counter()
    for pt in page_texts:
        for linha in {l.strip() for l in pt.splitlines() if l.strip()}:
            freq[linha] += 1
    boiler = {linha for linha, c in freq.items() if c >= threshold}
    if not boiler:
        return page_texts, []

    ordenadas: list[str] = []
    vistas: set[str] = set()
    limpos: list[str] = []
    for pt in page_texts:
        mantidas: list[str] = []
        for raw in pt.splitlines():
            if raw.strip() in boiler:
                if raw.strip() not in vistas:
                    vistas.add(raw.strip())
                    ordenadas.append(raw.strip())
                continue
            mantidas.append(raw)
        limpos.append("\n".join(mantidas))
    return limpos, ordenadas


def _tabela_para_markdown(tabela: list[list[Optional[str]]]) -> str:
    """Converte uma tabela (lista de linhas × células) em markdown GFM.

    Preserva a relação coluna↔valor que `extract_text()` perde quando
    achata a tabela em texto corrido. Células vazias viram `-`.
    Primeira linha vira cabeçalho (independente de ser ou não cabeçalho
    real — dá ao LLM um eixo para identificar colunas).
    """
    if not tabela or not tabela[0]:
        return ""

    n_cols = max(len(row) for row in tabela)

    def _normaliza(row: list[Optional[str]]) -> list[str]:
        cells = [
            _consertar_numeros_rachados((c or "").strip().replace("\n", " ").replace("|", "/"))
            for c in row
        ]
        cells += ["-"] * (n_cols - len(cells))
        return [c if c else "-" for c in cells]

    linhas = [_normaliza(row) for row in tabela]
    header = linhas[0]
    sep = ["---"] * n_cols
    corpo = linhas[1:]

    out = ["| " + " | ".join(header) + " |",
           "| " + " | ".join(sep) + " |"]
    for row in corpo:
        out.append("| " + " | ".join(row) + " |")
    return "\n".join(out)


def _ocr_page(page) -> str:
    """Roda Tesseract (pt) numa única página renderizada. Falha silenciosa → ''."""
    try:
        img = page.to_image(resolution=200).original
        return pytesseract.image_to_string(img, lang="por")
    except Exception as exc:
        logger.warning("OCR falhou na página: %s", exc)
        return ""


def extract_pdf(file_bytes: bytes) -> Tuple[str, str]:
    """Extract text from a PDF, preservando estrutura de tabelas.

    Para cada página: tenta `extract_tables()` primeiro. Quando há tabela,
    renderiza como markdown (preserva colunas semânticas — crítico para
    planilhas de aditivo onde "VALOR CONTRATO" vs "VALOR REPLANILHADO"
    ficam em colunas distintas). Texto não-tabular da mesma página é
    capturado por `extract_text()` em complemento.

    OCR POR PÁGINA (item 3.2 — caso Galícia): quando uma página não tem camada
    de texto mas contém imagem embutida (página digitalizada), roda Tesseract só
    naquela página. Antes o OCR só disparava quando o documento INTEIRO era escasso
    (média de chars/página baixa), então páginas escaneadas isoladas dentro de um PDF
    majoritariamente textual eram descartadas silenciosamente — e o modelo concluía
    pela ausência do documento que estava nos autos.

    Returns (text, method) where method is 'pdfplumber' or 'tesseract_ocr'
    (este último quando ao menos uma página precisou de OCR).
    """
    page_texts: list[str] = []
    table_blocos: list[str] = []
    total_chars = 0
    ocr_used = False

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        pages = pdf.pages
        if not pages:
            return "", "pdfplumber"

        for pidx, page in enumerate(pages):
            tabelas = []
            try:
                tabelas = page.extract_tables() or []
            except Exception as exc:
                logger.debug("extract_tables falhou em pág %d: %s", pidx + 1, exc)

            if tabelas:
                partes = [f"\n[Página {pidx + 1}]"]
                for tidx, tab in enumerate(tabelas):
                    md = _tabela_para_markdown(tab)
                    if md:
                        partes.append(f"\n**Tabela {tidx + 1}:**\n{md}")
                table_blocos.append("\n".join(partes))

            texto = page.extract_text() or ""
            if texto.strip():
                texto = _consertar_numeros_rachados(texto)
                page_texts.append(texto)
                total_chars += len(texto)
            elif page.images:
                # Página digitalizada (sem texto, com imagem): OCR só desta página.
                ocr = _ocr_page(page)
                ocr_used = True
                if ocr.strip():
                    page_texts.append(
                        f"[Página {pidx + 1} — documento digitalizado, lido por OCR]\n{ocr}"
                    )
                    total_chars += len(ocr)
                else:
                    # OCR não rendeu texto: registra a PRESENÇA do documento para o
                    # modelo não afirmar ausência (o relatório aceita "sinalizar dúvida").
                    page_texts.append(
                        f"[Página {pidx + 1} — documento digitalizado presente nos autos, "
                        "ilegível por OCR; NÃO afirmar ausência deste documento]"
                    )

    # Fallback extremo: nenhuma página tinha texto NEM imagem detectável (scan atípico
    # sem page.images) — OCR do documento inteiro, renderizando cada página.
    if total_chars == 0 and not ocr_used:
        logger.info("PDF sem texto e sem imagens detectáveis — OCR do documento inteiro")
        ocr_texts: list[str] = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                ocr_texts.append(_ocr_page(page))
        return "\n".join(ocr_texts), "tesseract_ocr"

    resumo = resumo_financeiro_canonico("\n".join(page_texts))
    corpo_texts, boilerplate = _separar_cabecalho_repetido(page_texts)

    blocos: list[str] = []
    if resumo:
        blocos.append(resumo)
    if boilerplate:
        blocos.append(
            "[CABEÇALHO/RODAPÉ REPETIDO DA PLANILHA]\n" + "\n".join(boilerplate)
        )
    blocos.extend(table_blocos)
    blocos.extend(corpo_texts)

    return "\n".join(blocos), ("tesseract_ocr" if ocr_used else "pdfplumber")


def extract_docx(file_bytes: bytes) -> Tuple[str, str]:
    """Extract text from a DOCX file (paragraphs and tables).

    Returns (text, 'python_docx').
    """
    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            parts.append(stripped)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts), "python_docx"


def _detect_encoding(file_bytes: bytes) -> str:
    result = chardet.detect(file_bytes)
    return result.get("encoding") or "utf-8"


def extract_file(
    filename: str, file_bytes: bytes
) -> Tuple[str, Optional[str], str]:
    """Dispatcher: extract text from a file based on its extension.

    Returns (text, method, status) where:
      - method is an ExtractionMethod enum value string or None
      - status is an ExtractionStatus enum value string: 'success', 'partial', 'failed'

    Handles password-protected files and encoding errors gracefully.
    """
    lower = filename.lower()
    try:
        if lower.endswith(".pdf"):
            text, method = extract_pdf(file_bytes)
            status = "success" if text.strip() else "partial"
            return _sanitize_text(text), method, status

        if lower.endswith(".docx"):
            text, method = extract_docx(file_bytes)
            status = "success" if text.strip() else "partial"
            return _sanitize_text(text), method, status

        if lower.endswith((".txt", ".text")):
            encoding = _detect_encoding(file_bytes)
            text = file_bytes.decode(encoding, errors="replace")
            return _sanitize_text(text), None, "success"

        if lower.endswith(".doc"):
            try:
                text = _extract_doc_via_libreoffice(file_bytes)
            except Exception as exc:
                logger.error("LibreOffice failed for %s: %s", filename, exc)
                return "", "fallback_libreoffice", "failed"
            status = "success" if text.strip() else "partial"
            return _sanitize_text(text), "fallback_libreoffice", status

        # Extension unknown or missing — try to detect via magic bytes
        kind = filetype.guess(file_bytes)
        if kind is not None:
            logger.info(
                "Detected file type by magic bytes: %s -> %s (%s)",
                filename, kind.extension, kind.mime,
            )
            guessed_name = f"_detected_.{kind.extension}"
            return extract_file(guessed_name, file_bytes)

        logger.warning("Unsupported file type for extraction: %s", filename)
        return "", None, "failed"

    except Exception as exc:
        msg = str(exc).lower()
        if "password" in msg or "encrypted" in msg:
            logger.warning("Password-protected or encrypted file skipped: %s", filename)
        else:
            logger.error("Extraction failed for %s: %s", filename, exc)
        return "", None, "failed"
