"""
docx_generator.py — Gerador programático determinístico de pareceres.

Portado das skills `parecer-lei-14133` (v1.4.0) e `parecer-municipal-geral`
(v1.4.0-pmg2). Reúne as duas em um único módulo: a vertente é informada
pela minuta e habilita ambos os tipos de marcador de revisão humana.

Diferenças em relação às skills:
- Não roda subprocess do auditor mecânico. O gate IRR-1/IRR-2 já é aplicado
  upstream em `parecer_engine._enforce_mechanical_gate` (Camada 4) sobre o
  `sections` do P2/P3, antes do conteúdo ser salvo como ParecerVersion.
  O .docx aqui é renderização determinística de conteúdo já gateado.
- Retorna `bytes` (BytesIO) em vez de salvar em disco — FastAPI faz streaming
  direto.
- Suporta marcadores das duas vertentes:
    - `[REVISAR — TEXTO]` (vertente licitação, IRR-4 da skill)
    - `[!VERIFICAR: TEXTO !]` (vertente municipal-geral, Regra ZT-5)
  Ambos renderizados em vermelho institucional RGB(192,0,0), negrito, inline.

Constantes de formatação byte-idênticas às skills (NÃO ALTERAR sem aprovação
do Dr. Ione):
- Margens A4: 2.5/3.0/3.0/3.0 cm
- Fonte corpo: Consolas 12pt, 1.5 entrelinhas
- Fonte cabeçalho: Garamond Small Caps
- Recuo primeira linha: 4.0 cm
- Ementa: figure-dash U+2015 entre palavras-chave
- Assinaturas: Consolas 10pt negrito, espaços manuais calibrados
  (25 Ione, 13 entre Matheus/Flávio, 26 entre OABs, 26 Valéria)
"""
from __future__ import annotations

import io
import re
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# =============================================================================
# CONSTANTES DE FORMATAÇÃO DO ESCRITÓRIO — não alterar sem aprovação do Dr. Ione
# =============================================================================

# Página A4
MARGEM_TOPO = Cm(2.5)
MARGEM_BASE = Cm(3.0)
MARGEM_ESQUERDA = Cm(3.0)
MARGEM_DIREITA = Cm(3.0)
HEADER_DISTANCE = Cm(1.25)
FOOTER_DISTANCE = Cm(1.25)

# Largura útil = 21 - 3 - 3 = 15 cm
LARGURA_UTIL_CM = 15.0

# Fontes
FONTE_CORPO = "Consolas"
FONTE_CABECALHO = "Garamond"
TAMANHO_CORPO = Pt(12)
TAMANHO_RODAPE = Pt(10)
TAMANHO_ASSINATURA = Pt(10)

# Recuo de primeira linha do corpo
RECUO_PRIMEIRA_LINHA = Cm(4.0)

# Espaçamento entre linhas
ESPACAMENTO_LINHA_15 = 1.5
ESPACAMENTO_LINHA_SIMPLES = 1.0

# Rodapé institucional
TEXTO_RODAPE = (
    "Rua Gen. Caiado de Castro 462, Luciano Cavalcante, Fortaleza-Ce, "
    "Fone (85) 3021-7701/ (85) 99981-4392/ (85) 99223-6716. "
    "Email: dr.ione@uol.com.br. Site: http://www.ioneadvogados.com.br"
)

# Bloco fixo de assinaturas
ASSINATURA_IONE = ("FRANCISCO IONE PEREIRA LIMA", "OAB-CE nº 4.585")
ASSINATURA_MATHEUS = ("MATHEUS NOGUEIRA PEREIRA LIMA", "OAB-CE nº 31.251")
ASSINATURA_FLAVIO = ("FLÁVIO HENRIQUE LUNA SILVA", "OAB-CE nº 31.252")
ASSINATURA_VALERIA = ("VALÉRIA MATIAS DE ALENCAR", "OAB/CE nº 36.666")


# =============================================================================
# MARCADORES DE REVISÃO HUMANA
# =============================================================================

# Cor institucional vermelha: #C00000.
COR_MARCADOR_REVISAO = RGBColor(0xC0, 0x00, 0x00)

# [REVISAR — TEXTO] (vertente licitação — IRR-4 da parecer-lei-14133)
PADRAO_MARCADOR_REVISAR = re.compile(
    r"\[REVISAR\s*[\-–—]\s*[^\]]+\]",
    re.IGNORECASE,
)

# [!VERIFICAR: TEXTO !] (vertente municipal-geral — Regra ZT-5)
PADRAO_MARCADOR_VERIFICAR = re.compile(
    r"\[!VERIFICAR:[^!]+!\]",
    re.IGNORECASE,
)

# União: usado para encontrar QUALQUER marcador (independente da vertente).
PADRAO_MARCADOR_QUALQUER = re.compile(
    r"\[REVISAR\s*[\-–—]\s*[^\]]+\]|\[!VERIFICAR:[^!]+!\]",
    re.IGNORECASE,
)


def contar_marcadores(texto) -> int:
    """Conta o número de marcadores [REVISAR—] ou [!VERIFICAR:!] no texto.

    Aceita str ou bloco {'text': ...} (parágrafo com recuos da régua)."""
    if isinstance(texto, dict):
        texto = texto.get("text", "")
    return len(PADRAO_MARCADOR_QUALQUER.findall(texto or ""))


def _coerce_block(item) -> Tuple[str, Optional[dict]]:
    """Normaliza um item de parágrafo em (texto, attrs).

    Aceita tanto `str` (gabarito antigo / placeholders) quanto bloco
    `{'text': str, 'attrs': {...}}` produzido por `minuta_from_tiptap`. `attrs`
    traz os recuos por parágrafo definidos na régua do editor (em cm) ou None."""
    if isinstance(item, dict):
        return item.get("text", ""), item.get("attrs")
    return item, None


def _aplicar_recuos_override(pf, attrs: Optional[dict], recuar_primeira_linha: bool) -> None:
    """Aplica recuos por parágrafo definidos na régua (cm), sobrepondo o padrão.

    Quando um atributo vem None, mantém o gabarito da casa: 1ª linha = 4 cm
    (se `recuar_primeira_linha`), esquerda/direita não são setadas (= 0). Assim,
    parágrafos intocados na régua saem idênticos ao formato calibrado do Dr. Ione."""
    attrs = attrs or {}
    fli = attrs.get("first_line_indent")
    if fli is not None:
        pf.first_line_indent = Cm(fli)
    elif recuar_primeira_linha:
        pf.first_line_indent = RECUO_PRIMEIRA_LINHA
    li = attrs.get("left_indent")
    if li is not None:
        pf.left_indent = Cm(li)
    ri = attrs.get("right_indent")
    if ri is not None:
        pf.right_indent = Cm(ri)


# =============================================================================
# HELPERS DE FORMATAÇÃO
# =============================================================================

def _set_font(
    run,
    font_name: str = FONTE_CORPO,
    size=TAMANHO_CORPO,
    bold: bool = False,
    italic: bool = False,
    small_caps: bool = False,
    color: Optional[RGBColor] = None,
) -> None:
    """Aplica fonte a um run padronizadamente, também no XML para compatibilidade."""
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if small_caps:
        run.font.small_caps = True
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def _adicionar_runs_com_marcadores(
    paragrafo,
    texto: str,
    bold: bool = False,
    fonte_size=TAMANHO_CORPO,
) -> None:
    """Adiciona runs detectando e renderizando marcadores [REVISAR—] / [!VERIFICAR:!]
    em VERMELHO/NEGRITO inline. Sem marcadores, comporta-se como _set_font padrão."""
    matches = list(PADRAO_MARCADOR_QUALQUER.finditer(texto))

    if not matches:
        run = paragrafo.add_run(texto)
        _set_font(run, bold=bold, size=fonte_size)
        return

    cursor = 0
    for m in matches:
        inicio, fim = m.span()
        if inicio > cursor:
            run_normal = paragrafo.add_run(texto[cursor:inicio])
            _set_font(run_normal, bold=bold, size=fonte_size)

        run_marcador = paragrafo.add_run(texto[inicio:fim])
        _set_font(
            run_marcador,
            bold=True,
            size=fonte_size,
            color=COR_MARCADOR_REVISAO,
        )
        cursor = fim

    if cursor < len(texto):
        run_final = paragrafo.add_run(texto[cursor:])
        _set_font(run_final, bold=bold, size=fonte_size)


def _adicionar_paragrafo_corpo(
    doc,
    texto: str,
    alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY,
    recuar_primeira_linha: bool = True,
    negrito: bool = False,
    fonte_size=TAMANHO_CORPO,
    attrs: Optional[dict] = None,
):
    """Parágrafo de corpo com formatação padrão e detecção automática de marcadores.
    space_after = 6pt — alinhado ao gabarito do escritório (cliente).

    `attrs` (opcional) traz recuos por parágrafo definidos na régua do editor
    (first_line_indent / left_indent / right_indent em cm); None = padrão da casa."""
    p = doc.add_paragraph()
    p.alignment = alinhamento
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = ESPACAMENTO_LINHA_15
    _aplicar_recuos_override(pf, attrs, recuar_primeira_linha)
    _adicionar_runs_com_marcadores(p, texto, bold=negrito, fonte_size=fonte_size)
    return p


def _adicionar_blockquote(doc, texto: str):
    """Citação legal em bloco — padrão calibrado pelo Dr. Ione.

    align=justify, left_indent=4cm (parágrafo inteiro recuado, não só 1ª linha),
    Consolas 12pt **negrito**, espaço de 12pt antes e depois para destacar
    visualmente do corpo. Sem prefixo '>' nem itálico via '*'.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = RECUO_PRIMEIRA_LINHA
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    pf.line_spacing = ESPACAMENTO_LINHA_15
    _adicionar_runs_com_marcadores(p, texto, bold=True)
    return p


def _adicionar_paragrafo_vazio(doc):
    p = doc.add_paragraph()
    run = p.add_run("")
    _set_font(run)
    return p


def _inserir_campo_page(run) -> None:
    """Insere o campo PAGE (número da página dinâmico) em um run."""
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)


# =============================================================================
# CONSTRUÇÃO DA PÁGINA
# =============================================================================

def _configurar_pagina(doc) -> None:
    section = doc.sections[0]
    section.top_margin = MARGEM_TOPO
    section.bottom_margin = MARGEM_BASE
    section.left_margin = MARGEM_ESQUERDA
    section.right_margin = MARGEM_DIREITA
    section.header_distance = HEADER_DISTANCE
    section.footer_distance = FOOTER_DISTANCE


def _construir_cabecalho(doc) -> None:
    """Cabeçalho: linha 1 'Advocacia & Assessoria' Garamond Small Caps centralizado;
    linha 2 'Dr. Francisco Ione Pereira Lima' Garamond Small Caps negrito centralizado;
    linha 3 número da página em campo PAGE alinhado à direita (parágrafo separado
    para preservar centralização da linha 2)."""
    header = doc.sections[0].header

    p1 = header.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run("Advocacia & Assessoria")
    _set_font(run1, font_name=FONTE_CABECALHO, small_caps=True)

    p2 = header.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_nome = p2.add_run("Dr. Francisco Ione Pereira Lima")
    _set_font(run_nome, font_name=FONTE_CABECALHO, bold=True, small_caps=True)

    p3 = header.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_page = p3.add_run()
    _set_font(run_page, font_name=FONTE_CABECALHO, bold=True)
    _inserir_campo_page(run_page)


def _construir_rodape(doc) -> None:
    """Rodapé em uma linha, Consolas 10pt, alinhado à esquerda."""
    footer = doc.sections[0].footer
    pf = footer.paragraphs[0]
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.paragraph_format.space_before = Pt(0)
    pf.paragraph_format.space_after = Pt(0)
    pf.paragraph_format.line_spacing = ESPACAMENTO_LINHA_SIMPLES

    run = pf.add_run(TEXTO_RODAPE)
    _set_font(run, size=TAMANHO_RODAPE)


# =============================================================================
# BLOCOS DO CORPO
# =============================================================================

def _adicionar_titulo_parecer(doc) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("PARECER JURÍDICO")
    _set_font(run, bold=True)


def _adicionar_consulente(doc, consulente_texto: str) -> None:
    """'CONSULENTE: [texto]' — apenas o rótulo em negrito, parágrafo justificado."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)

    run_label = p.add_run("CONSULENTE: ")
    _set_font(run_label, bold=True)

    run_valor = p.add_run(consulente_texto)
    _set_font(run_valor, bold=False)


def _adicionar_ementa(doc, palavras_chave: List[str]) -> None:
    """Ementa em palavras-chave separadas por figure-dash (U+2015), MAIÚSCULAS, negrito."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = RECUO_PRIMEIRA_LINHA
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    pf.line_spacing = ESPACAMENTO_LINHA_15

    figure_dash = "―"
    texto_ementa = "EMENTA: " + (" " + figure_dash + " ").join(palavras_chave) + "."
    texto_ementa = texto_ementa.upper()  # IRR-1

    run = p.add_run(texto_ementa)
    _set_font(run, bold=True)


def _adicionar_titulo_secao(doc, titulo: str) -> None:
    """Títulos das três seções: 'I — RELATÓRIO', 'II — FUNDAMENTOS', 'III — CONCLUSÃO'."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(titulo)
    _set_font(run, bold=True)


_PADRAO_BLOCKQUOTE_LITERAL = re.compile(
    r"^>\s*\*?\s*[\"'“”]?(.*?)[\"'“”]?\s*\*?\s*$",
    re.DOTALL,
)


def _normalizar_blockquote(texto: str) -> str:
    """Remove prefixo '>' (markdown vazado do TipTap) e wrap '*...*' (itálico
    markdown) do texto de citação legal, devolvendo apenas o conteúdo."""
    linhas = [linha.lstrip("> ").strip() for linha in texto.splitlines()]
    bruto = " ".join(linhas).strip()
    # Strip de wrap '*...*' (negrito/itálico markdown) e aspas curvas
    m = re.match(r"^\*+\s*(.*?)\s*\*+$", bruto, re.DOTALL)
    if m:
        bruto = m.group(1).strip()
    return bruto


def _adicionar_paragrafos_corpo(doc, paragrafos: List) -> None:
    """Renderiza parágrafos de uma seção. Cada item pode ser `str` ou bloco
    `{'text','attrs'}`. Linhas que começam com '>' viram citação legal
    (blockquote) com formatação calibrada fixa — ignoram os recuos da régua."""
    for item in paragrafos:
        texto, attrs = _coerce_block(item)
        if texto.lstrip().startswith(">"):
            _adicionar_blockquote(doc, _normalizar_blockquote(texto))
        else:
            _adicionar_paragrafo_corpo(doc, texto, attrs=attrs)


def _adicionar_recomendacoes(doc, alineas: List[Tuple[str, str]]) -> None:
    """Alíneas (a), (b), (c), (d) das recomendações — formatadas como corpo normal
    (align=both, fli=4cm, space_after=6pt), conforme gabarito do escritório."""
    for letra, texto in alineas:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.first_line_indent = RECUO_PRIMEIRA_LINHA
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing = ESPACAMENTO_LINHA_15

        run_letra = p.add_run(f"({letra}) ")
        _set_font(run_letra)
        _adicionar_runs_com_marcadores(p, texto, bold=False)


# =============================================================================
# BLOCO DE ASSINATURAS — espaços manuais calibrados pelo escritório
# =============================================================================

def _adicionar_fortaleza_data(doc, data_extenso: str) -> None:
    """Linha 'Fortaleza/CE, [data].' — formatada como parágrafo de corpo (align=both,
    fli=4cm), com espaço extra antes e depois para separar do bloco de assinaturas.
    Alinhado ao gabarito do escritório."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = RECUO_PRIMEIRA_LINHA
    pf.space_before = Pt(24)
    pf.space_after = Pt(24)
    pf.line_spacing = ESPACAMENTO_LINHA_15
    run = p.add_run(data_extenso + ".")
    _set_font(run)


def _adicionar_paragrafo_assinatura_centralizado(
    doc, nome: str, oab: str, espacos: int = 25
) -> None:
    """Bloco de assinatura no padrão real do escritório:
    align=LEFT, mesmo número de espaços de avanço para nome e OAB,
    Consolas 10pt negrito. Garante alinhamento mecânico nome/OAB."""
    avanco = " " * espacos

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf1 = p1.paragraph_format
    pf1.space_before = Pt(0)
    pf1.space_after = Pt(0)
    pf1.line_spacing = ESPACAMENTO_LINHA_SIMPLES

    run_esp1 = p1.add_run(avanco)
    _set_font(run_esp1, bold=True, size=TAMANHO_ASSINATURA)
    run_nome = p1.add_run(nome)
    _set_font(run_nome, bold=True, size=TAMANHO_ASSINATURA)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf2 = p2.paragraph_format
    pf2.space_before = Pt(0)
    pf2.space_after = Pt(0)
    pf2.line_spacing = ESPACAMENTO_LINHA_SIMPLES

    run_esp2 = p2.add_run(avanco)
    _set_font(run_esp2, bold=True, size=TAMANHO_ASSINATURA)
    run_oab = p2.add_run(oab)
    _set_font(run_oab, bold=True, size=TAMANHO_ASSINATURA)


def _adicionar_assinatura_dupla_matheus_flavio(doc) -> None:
    """Bloco Matheus + Flávio: parágrafo único por linha, espaços manuais calibrados.
    Linha nomes: 29 chars MATHEUS + 13 espaços + 26 chars FLÁVIO.
    Linha OABs:  16 chars OAB-M + 26 espaços + 16 chars OAB-F.
    Consolas 10pt negrito cabe em 15 cm úteis."""
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf1 = p1.paragraph_format
    pf1.space_before = Pt(0)
    pf1.space_after = Pt(0)
    pf1.line_spacing = ESPACAMENTO_LINHA_SIMPLES

    run_m = p1.add_run("MATHEUS NOGUEIRA PEREIRA LIMA")
    _set_font(run_m, bold=True, size=TAMANHO_ASSINATURA)
    run_esp1 = p1.add_run(" " * 13)
    _set_font(run_esp1, bold=True, size=TAMANHO_ASSINATURA)
    run_f = p1.add_run("FLÁVIO HENRIQUE LUNA SILVA")
    _set_font(run_f, bold=True, size=TAMANHO_ASSINATURA)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf2 = p2.paragraph_format
    pf2.space_before = Pt(0)
    pf2.space_after = Pt(0)
    pf2.line_spacing = ESPACAMENTO_LINHA_SIMPLES

    run_om = p2.add_run("OAB-CE nº 31.251")
    _set_font(run_om, bold=True, size=TAMANHO_ASSINATURA)
    run_esp2 = p2.add_run(" " * 26)
    _set_font(run_esp2, bold=True, size=TAMANHO_ASSINATURA)
    run_of = p2.add_run("OAB-CE nº 31.252")
    _set_font(run_of, bold=True, size=TAMANHO_ASSINATURA)


def _adicionar_bloco_assinaturas(doc) -> None:
    """Três blocos: Ione (25 espaços) — Matheus+Flávio (parágrafo único) — Valéria (26 espaços)."""
    _adicionar_paragrafo_assinatura_centralizado(doc, *ASSINATURA_IONE, espacos=25)
    _adicionar_paragrafo_vazio(doc)
    _adicionar_assinatura_dupla_matheus_flavio(doc)
    _adicionar_paragrafo_vazio(doc)
    _adicionar_paragrafo_assinatura_centralizado(doc, *ASSINATURA_VALERIA, espacos=26)


def _adicionar_fecho(doc) -> None:
    """'É o parecer, submetido à superior consideração.' — space_after = 6pt
    para alinhar com o restante dos parágrafos de corpo (gabarito do escritório)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = RECUO_PRIMEIRA_LINHA
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = ESPACAMENTO_LINHA_15
    run = p.add_run("É o parecer, submetido à superior consideração.")
    _set_font(run)


# =============================================================================
# FUNÇÃO PÚBLICA — ÚNICA INTERFACE DO MÓDULO
# =============================================================================

CAMPOS_OBRIGATORIOS = (
    "consulente",
    "ementa_palavras_chave",
    "relatorio_paragrafos",
    "fundamentos_paragrafos",
    "conclusao_dispositivo",
    "recomendacoes_alineas",
    "data_extenso",
)


def gerar_parecer_bytes(minuta: dict) -> bytes:
    """Renderiza um parecer .docx no padrão do escritório a partir da minuta.

    Args:
        minuta: dicionário com:
            - consulente (str)
            - ementa_palavras_chave (List[str]) — última deve ser a conclusão
              ("PARECER FAVORÁVEL", "PARECER PERTINENTE — PROVIMENTO RECOMENDADO" etc.)
            - relatorio_paragrafos (List[str]) — último é "É o breve relatório.
              Passa-se à fundamentação."
            - fundamentos_paragrafos (List[str]) — prosa contínua, sem subdivisão.
              Parágrafos iniciados por '>' viram blockquote (citação legal recuada).
            - conclusao_dispositivo (str) — "Diante do exposto, ..."
            - recomendacoes_alineas (List[Tuple[str, str]]) — [("a", "..."), ...]
            - data_extenso (str) — "Fortaleza/CE, 15 de maio de 2026"
            - subtipo (str) — informativo
            - vertente (str) — informativo

    Returns:
        bytes do arquivo .docx pronto para download.

    Raises:
        KeyError: se um campo obrigatório estiver ausente.
    """
    faltantes = [campo for campo in CAMPOS_OBRIGATORIOS if campo not in minuta]
    if faltantes:
        raise KeyError(
            f"Campos obrigatórios ausentes na minuta: {faltantes}. "
            f"Obrigatórios: {list(CAMPOS_OBRIGATORIOS)}"
        )

    doc = Document()

    _configurar_pagina(doc)
    _construir_cabecalho(doc)
    _construir_rodape(doc)

    _adicionar_titulo_parecer(doc)
    _adicionar_consulente(doc, minuta["consulente"])
    _adicionar_ementa(doc, minuta["ementa_palavras_chave"])

    _adicionar_titulo_secao(doc, "I — RELATÓRIO")
    _adicionar_paragrafos_corpo(doc, minuta["relatorio_paragrafos"])

    _adicionar_titulo_secao(doc, "II — FUNDAMENTOS")
    _adicionar_paragrafos_corpo(doc, minuta["fundamentos_paragrafos"])

    _adicionar_titulo_secao(doc, "III — CONCLUSÃO")
    _disp_texto, _disp_attrs = _coerce_block(minuta["conclusao_dispositivo"])
    _adicionar_paragrafo_corpo(doc, _disp_texto, attrs=_disp_attrs)
    _adicionar_recomendacoes(doc, minuta["recomendacoes_alineas"])

    _adicionar_fecho(doc)
    _adicionar_fortaleza_data(doc, minuta["data_extenso"])
    _adicionar_bloco_assinaturas(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# =============================================================================
# PARSER TipTap JSON → minuta dict
# =============================================================================

# Variantes de hífen/travessão tolerados nos títulos de seção
_SECTION_TITLES = {
    "ementa": ("EMENTA",),
    "relatorio": ("I — RELATÓRIO", "I – RELATÓRIO", "I - RELATÓRIO", "RELATÓRIO"),
    "fundamentos": (
        "II — FUNDAMENTOS",
        "II – FUNDAMENTOS",
        "II - FUNDAMENTOS",
        "FUNDAMENTOS",
    ),
    "conclusao": (
        "III — CONCLUSÃO",
        "III – CONCLUSÃO",
        "III - CONCLUSÃO",
        "CONCLUSÃO",
    ),
}

# Alínea no início de parágrafo: "a)", "(a)", "a-)" etc.
_PADRAO_ALINEA = re.compile(r"^\(?([a-z])\)\s*(.*)$", re.DOTALL)

# Separadores de palavras-chave da ementa: figure-dash (U+2015) ou em-dash
_PADRAO_SEP_EMENTA = re.compile(r"\s*[―—–]\s*")


def _tiptap_node_to_text(node: dict) -> str:
    """Concatena o texto plano de um node TipTap (paragraph, heading, blockquote, list).

    Blockquote é marcado com prefixo '>' no início de cada linha — esse sentinel
    é consumido por `_adicionar_paragrafos_corpo` para despachar ao render de
    citação legal (left_indent=4cm, bold, space_before/after=12pt).
    """
    if node.get("type") == "text":
        return node.get("text", "")
    parts = []
    for child in node.get("content", []):
        parts.append(_tiptap_node_to_text(child))
    inner = "".join(parts)
    if node.get("type") == "blockquote":
        lines = inner.split("\n")
        return "\n".join(f"> {line}" for line in lines)
    return inner


def _extract_indent_attrs(node: dict) -> dict:
    """Lê os recuos por parágrafo (em cm) gravados pela régua do editor.

    A extensão frontend `ParagraphIndent` grava `firstLineIndent`/`leftIndent`/
    `rightIndent` em `node.attrs` (None = padrão da casa). Só `paragraph` os carrega;
    headings e blockquotes não têm os atributos, logo seguem o formato calibrado."""
    a = node.get("attrs") or {}
    return {
        "first_line_indent": a.get("firstLineIndent"),
        "left_indent": a.get("leftIndent"),
        "right_indent": a.get("rightIndent"),
    }


def _tiptap_node_to_block(node: dict) -> dict:
    """Converte um node TipTap em bloco `{'text', 'attrs'}` carregando os recuos."""
    return {"text": _tiptap_node_to_text(node), "attrs": _extract_indent_attrs(node)}


def _identify_section(heading_text: str) -> Optional[str]:
    """Mapeia o texto de um heading H2 para a chave canônica da seção."""
    cleaned = heading_text.strip().upper()
    for key, variants in _SECTION_TITLES.items():
        for variant in variants:
            if variant.upper() in cleaned:
                return key
    return None


def _group_sections(tiptap: dict) -> dict:
    """Percorre os nodes TipTap e devolve {ementa: [blocks], relatorio: [...], ...}.

    Cada bloco é `{'text', 'attrs'}` — `attrs` carrega os recuos por parágrafo da
    régua do editor (ver `_extract_indent_attrs`)."""
    grouped: dict = {key: [] for key in _SECTION_TITLES.keys()}
    current: Optional[str] = None
    for node in tiptap.get("content", []):
        ntype = node.get("type", "")
        if ntype == "heading" and node.get("attrs", {}).get("level") == 2:
            heading_text = "".join(
                c.get("text", "") for c in node.get("content", []) if c.get("type") == "text"
            )
            mapped = _identify_section(heading_text)
            current = mapped
            continue
        if current is None:
            continue
        block = _tiptap_node_to_block(node)
        block["text"] = block["text"].strip()
        if block["text"]:
            grouped[current].append(block)
    return grouped


def _parse_ementa_palavras_chave(blocos: List) -> List[str]:
    """Extrai as palavras-chave da ementa (só-texto; recuos da régua não se aplicam).

    O texto da ementa pode chegar como:
    - "EMENTA: PALAVRA1 ― PALAVRA2 ― PARECER FAVORÁVEL."
    - "PALAVRA1 ― PALAVRA2 ― PARECER FAVORÁVEL"
    - vários parágrafos (raro — junta tudo em uma linha)
    """
    if not blocos:
        return []
    bruto = " ".join(_coerce_block(b)[0] for b in blocos).strip()
    bruto = re.sub(r"^EMENTA\s*:\s*", "", bruto, flags=re.IGNORECASE)
    bruto = bruto.rstrip(".")
    palavras = [p.strip() for p in _PADRAO_SEP_EMENTA.split(bruto) if p.strip()]
    return palavras


def _parse_conclusao(blocos: List) -> Tuple[dict, List[Tuple[str, str]]]:
    """Quebra a conclusão em (dispositivo, alíneas).

    - Dispositivo = bloco `{'text','attrs'}` com os parágrafos antes da primeira
      alínea concatenados ("\\n\\n"). Os recuos da régua vêm do 1º desses parágrafos.
    - Alíneas = parágrafos que começam com "a)", "(b)", etc. (só-texto, padrão da casa).

    Parágrafos posteriores às alíneas (eventuais "Cumpre advertir..." que escapem
    do filtro do prompt) são silenciosamente descartados — o formato de
    advertência terminal ao gestor foi abolido pelo Dr. Ione.
    """
    dispositivo_parts: List[str] = []
    dispositivo_attrs: Optional[dict] = None
    alineas: List[Tuple[str, str]] = []
    estado = "antes"  # "antes" → "alineas" → "descartado"

    for bloco in blocos:
        texto, attrs = _coerce_block(bloco)
        match = _PADRAO_ALINEA.match(texto.strip())
        if match:
            letra = match.group(1).lower()
            conteudo = match.group(2).strip()
            alineas.append((letra, conteudo))
            estado = "alineas"
            continue
        if estado == "antes":
            if dispositivo_attrs is None:
                dispositivo_attrs = attrs
            dispositivo_parts.append(texto)
        # else: parágrafo pós-alíneas — descartado

    dispositivo = {
        "text": "\n\n".join(dispositivo_parts).strip(),
        "attrs": dispositivo_attrs,
    }
    return dispositivo, alineas


def _ensure_relatorio_termina_com_passa_se(paragrafos: List) -> List:
    """Garante que o último parágrafo do relatório é a fórmula de passagem.

    Opera sobre blocos `{'text','attrs'}`; a fórmula anexada usa o padrão da casa
    (attrs None)."""
    if not paragrafos:
        return paragrafos
    formula = "É o breve relatório. Passa-se à fundamentação."
    ultimo = _coerce_block(paragrafos[-1])[0].strip()
    if ultimo == formula or formula in ultimo:
        return paragrafos
    return paragrafos + [{"text": formula, "attrs": None}]


def minuta_from_tiptap(
    tiptap: dict,
    consulente: str,
    data_extenso: str,
    subtipo: str = "",
    vertente: str = "",
) -> dict:
    """Converte o TipTap JSON salvo em uma `minuta` dict consumível pelo gerador.

    Args:
        tiptap: dict TipTap (content_tiptap da ParecerVersion mais recente).
        consulente: texto do consulente já formatado (ex: "Pregoeiro do Município de...").
        data_extenso: "Fortaleza/CE, 15 de maio de 2026" — sem ponto final (o gerador adiciona).
        subtipo: rótulo informativo.
        vertente: rótulo informativo.

    Returns:
        dict no formato esperado por `gerar_parecer_bytes`.
    """
    grouped = _group_sections(tiptap or {})

    ementa_palavras = _parse_ementa_palavras_chave(grouped.get("ementa", []))
    relatorio = _ensure_relatorio_termina_com_passa_se(grouped.get("relatorio", []))
    fundamentos = grouped.get("fundamentos", [])
    dispositivo, alineas = _parse_conclusao(grouped.get("conclusao", []))
    # Colapsa para "" quando vazio — preserva a salvaguarda `if not ...` do export_service
    # (um dict sempre seria truthy e burlaria o placeholder de dispositivo ausente).
    conclusao_dispositivo = dispositivo if dispositivo["text"] else ""

    return {
        "consulente": consulente,
        "ementa_palavras_chave": ementa_palavras,
        "relatorio_paragrafos": relatorio,
        "fundamentos_paragrafos": fundamentos,
        "conclusao_dispositivo": conclusao_dispositivo,
        "recomendacoes_alineas": alineas,
        "data_extenso": data_extenso,
        "subtipo": subtipo,
        "vertente": vertente,
    }
