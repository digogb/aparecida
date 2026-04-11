"""
Export service: converts parecer content to DOCX and PDF formats.

Formato baseado no template profissional do escritório:
- Advocacia & Assessoria — Dr. Francisco Ione Pereira Lima
- Font: Times New Roman (serif), 12pt
- Seções: EMENTA, I — RELATÓRIO, II — FUNDAMENTOS, III — CONCLUSÃO
- Assinaturas: 2×2 grid com OAB
- Rodapé: endereço do escritório
"""
from __future__ import annotations

import io
import logging
from datetime import datetime, timezone
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.parecer import ParecerRequest, ParecerVersion

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Escritório info (fixo — dados do template de referência)
# ---------------------------------------------------------------------------
ESCRITORIO_LINHA1 = "Pearson Hardman"
ESCRITORIO_LINHA2 = "Advocacia & Assessoria"
ESCRITORIO_ENDERECO = "E-mail: contato@pearsonhardman.com"

ADVOGADOS = [
    ("Harvey Specter", "OAB/SP 12.345"),
    ("Mike Ross", "OAB/SP 23.456"),
    ("Louis Litt", "OAB/SP 34.567"),
]

MESES = [
    "", "janeiro", "fevereiro", "março", "abril", "maio", "junho",
    "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
]


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _formatar_data_extenso(municipio_uf: str = "Fortaleza/CE") -> str:
    now = datetime.now(timezone.utc)
    return f"{municipio_uf}, {now.day} de {MESES[now.month]} de {now.year}."


_ADVOGADO_NOMES = {n for n, _ in ADVOGADOS}
_SKIP_TEXTS = {
    ESCRITORIO_LINHA1, ESCRITORIO_LINHA2, "PARECER JURÍDICO",
}


def _is_decoration_node(node: dict[str, Any]) -> bool:
    """Detect nodes that are 'decoration' (header, signatures, footer) — skip in export."""
    text = _extract_text(node.get("content", []))
    if not text:
        return False
    # Header lines
    if text.strip() in _SKIP_TEXTS:
        return True
    # Identification fields (added by builder)
    if text.startswith("Órgão Consulente:") or text.startswith("Referência:") or text.startswith("Assunto:"):
        return True
    # Signature block: contains multiple OAB references
    if text.count("OAB/") >= 2:
        return True
    # Individual advogado name or OAB line
    stripped = text.strip()
    if stripped in _ADVOGADO_NOMES or stripped.startswith("OAB/"):
        return True
    # Footer (endereço)
    if "Rua Gen. Caiado de Castro" in text:
        return True
    # Date line (added by builder)
    for m in MESES[1:]:
        if f" de {m} de " in text and text.strip().endswith("."):
            return True
    return False


def _filter_tiptap_content(content: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Remove decoration nodes from TipTap content — header, signatures, footer."""
    return [node for node in content if not _is_decoration_node(node)]


def _get_metadata(req: ParecerRequest) -> dict:
    """Extract metadata from classificacao or fallback to request fields."""
    classification = req.classificacao or {}
    municipio = classification.get("municipio", "")
    uf = classification.get("uf", "CE")
    municipio_uf = f"{municipio}/{uf}" if municipio else "[Município/UF]"
    return {
        "orgao_consulente": classification.get("orgao_consulente", req.sender_email or "[Órgão não informado]"),
        "municipio_uf": municipio_uf,
        "assunto": classification.get("assunto_resumido", req.subject or "[Assunto]"),
        "referencia": req.numero_parecer or "N/A",
    }


def _add_paragraph_border_bottom(paragraph) -> None:
    """Add a bottom border to a paragraph via XML."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_border(cell, **kwargs):
    """Set borders on a table cell. kwargs: top, bottom, start, end with val, sz, color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, attrs in kwargs.items():
        element = OxmlElement(f"w:{edge}")
        for attr_name, attr_val in attrs.items():
            element.set(qn(f"w:{attr_name}"), str(attr_val))
        tcBorders.append(element)
    tcPr.append(tcBorders)


def _remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


# ---------------------------------------------------------------------------
# DOCX: Header
# ---------------------------------------------------------------------------

def _build_header(doc: Document) -> None:
    """Cabeçalho: Pearson Hardman / Advocacia & Assessoria."""
    # Linha 1
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.space_before = Pt(0)
    run1 = p1.add_run(ESCRITORIO_LINHA1)
    run1.font.name = "Times New Roman"
    run1.font.size = Pt(13)
    run1.font.small_caps = True

    # Linha 2
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.space_before = Pt(2)
    run2 = p2.add_run(ESCRITORIO_LINHA2)
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(11)
    run2.bold = True
    run2.font.small_caps = True

    # Bottom border on last header paragraph
    _add_paragraph_border_bottom(p2)


# ---------------------------------------------------------------------------
# DOCX: Title + Identification
# ---------------------------------------------------------------------------

def _build_title_and_meta(doc: Document, req: ParecerRequest) -> None:
    meta = _get_metadata(req)

    # Título
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(14)
    title_p.paragraph_format.space_after = Pt(14)
    run = title_p.add_run("PARECER JURÍDICO")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)

    # Identification fields
    fields = [
        ("Órgão Consulente:", meta["orgao_consulente"]),
        ("Referência:", meta["referencia"]),
        ("Assunto:", meta["assunto"]),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        run_label = p.add_run(label + " ")
        run_label.bold = True
        run_label.font.name = "Times New Roman"
        run_label.font.size = Pt(11.5)
        run_value = p.add_run(value)
        run_value.font.name = "Times New Roman"
        run_value.font.size = Pt(11.5)


# ---------------------------------------------------------------------------
# DOCX: TipTap JSON → python-docx mapping
# ---------------------------------------------------------------------------

def _add_tiptap_content(doc: Document, content: list[dict[str, Any]]) -> None:
    """Map TipTap JSON nodes to python-docx elements with professional formatting."""
    for node in content:
        node_type = node.get("type", "")

        if node_type == "paragraph":
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(6)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Text indent for normal paragraphs
            para.paragraph_format.first_line_indent = Cm(1.25)
            _add_inline_content(para, node.get("content", []))

        elif node_type == "heading":
            level = node.get("attrs", {}).get("level", 2)
            text = _extract_text(node.get("content", []))

            if level == 2:
                # Section heading (EMENTA, I — RELATÓRIO, etc.)
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(18)
                para.paragraph_format.space_after = Pt(8)
                run = para.add_run(text)
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            elif level == 3:
                # Numbered subtitle (1. Da Legislação, etc.)
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(4)
                run = para.add_run(text)
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            else:
                para = doc.add_paragraph()
                _add_inline_content(para, node.get("content", []))

        elif node_type == "blockquote":
            # Blockquote: italic, indented left
            for child in node.get("content", []):
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Cm(2)
                para.paragraph_format.right_indent = Cm(1)
                para.paragraph_format.space_after = Pt(4)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _add_inline_content(para, child.get("content", []), italic=True)

        elif node_type == "bulletList":
            for item in node.get("content", []):
                if item.get("type") == "listItem":
                    for child in item.get("content", []):
                        para = doc.add_paragraph(style="List Bullet")
                        _add_inline_content(para, child.get("content", []))

        elif node_type == "orderedList":
            for item in node.get("content", []):
                if item.get("type") == "listItem":
                    for child in item.get("content", []):
                        para = doc.add_paragraph(style="List Number")
                        _add_inline_content(para, child.get("content", []))

        else:
            children = node.get("content", [])
            if children:
                _add_tiptap_content(doc, children)


def _add_inline_content(
    paragraph, content: list[dict[str, Any]], italic: bool = False
) -> None:
    """Add inline text nodes (with marks) to a paragraph."""
    for item in content:
        if item.get("type") == "text":
            text = item.get("text", "")
            run = paragraph.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

            if italic:
                run.italic = True

            for mark in item.get("marks", []):
                mark_type = mark.get("type", "")
                if mark_type == "bold":
                    run.bold = True
                elif mark_type == "italic":
                    run.italic = True
                elif mark_type == "underline":
                    run.underline = True
                elif mark_type == "strike":
                    run.font.strike = True
        elif item.get("type") == "hardBreak":
            paragraph.add_run("\n")


def _extract_text(content: list[dict[str, Any]]) -> str:
    """Extract plain text from TipTap inline content."""
    return "".join(item.get("text", "") for item in content if item.get("type") == "text")


# ---------------------------------------------------------------------------
# DOCX: Local/Data + Assinaturas + Rodapé
# ---------------------------------------------------------------------------

def _build_closing(doc: Document, req: ParecerRequest) -> None:
    """Local, data e bloco de assinaturas."""
    meta = _get_metadata(req)

    # Local e data (right-aligned)
    data_p = doc.add_paragraph()
    data_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    data_p.paragraph_format.space_before = Pt(24)
    data_p.paragraph_format.space_after = Pt(6)
    run = data_p.add_run(_formatar_data_extenso(meta["municipio_uf"]))
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Assinaturas: 2×2 table sem bordas
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(table)

    for idx, (nome, oab) in enumerate(ADVOGADOS):
        row = idx // 2
        col = idx % 2
        cell = table.cell(row, col)

        # Clear default paragraph
        cell.paragraphs[0].clear()

        # Spacing before signature line
        spacer = cell.paragraphs[0]
        spacer.paragraph_format.space_before = Pt(36)
        spacer.paragraph_format.space_after = Pt(0)

        # Add top border to simulate signature line
        _set_cell_border(cell, top={"val": "single", "sz": "4", "color": "000000"})

        # Name (small caps, bold, centered)
        name_run = spacer.add_run(nome)
        name_run.bold = True
        name_run.font.name = "Times New Roman"
        name_run.font.size = Pt(11)
        name_run.font.small_caps = True
        spacer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # OAB
        oab_p = cell.add_paragraph()
        oab_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        oab_p.paragraph_format.space_before = Pt(2)
        oab_p.paragraph_format.space_after = Pt(20)
        oab_run = oab_p.add_run(oab)
        oab_run.font.name = "Times New Roman"
        oab_run.font.size = Pt(10)


def _build_footer(doc: Document) -> None:
    """Rodapé: endereço do escritório com borda superior."""
    doc.add_paragraph()  # spacing

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(10)

    # Top border
    pPr = footer_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "000000")
    pBdr.append(top)
    pPr.append(pBdr)

    run = footer_p.add_run(ESCRITORIO_ENDERECO)
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ---------------------------------------------------------------------------
# DOCX generation (public)
# ---------------------------------------------------------------------------

async def to_docx(
    parecer_request: ParecerRequest,
    version: ParecerVersion,
    db: AsyncSession,
) -> bytes:
    """Generate DOCX bytes from a parecer version's TipTap content."""
    doc = Document()

    # Page setup (A4)
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.6

    # Build document
    _build_header(doc)
    _build_title_and_meta(doc, parecer_request)

    # Body from TipTap JSON (filter out decoration: header, sigs, footer)
    tiptap = version.content_tiptap
    if tiptap and isinstance(tiptap, dict):
        content = _filter_tiptap_content(tiptap.get("content", []))
        _add_tiptap_content(doc, content)
    elif version.content_html:
        doc.add_paragraph(version.content_html)

    # Closing: local/date + signatures
    _build_closing(doc, parecer_request)

    # Footer
    _build_footer(doc)

    # Serialize
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ---------------------------------------------------------------------------
# PDF generation (public)
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# PDF: TipTap JSON → HTML (mesma estrutura do DOCX)
# ---------------------------------------------------------------------------

def _escape_html(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _tiptap_inline_to_html(content: list[dict[str, Any]], force_italic: bool = False) -> str:
    """Convert TipTap inline content (text nodes with marks) to HTML."""
    parts: list[str] = []
    for item in content:
        if item.get("type") == "text":
            text = _escape_html(item.get("text", ""))
            marks = item.get("marks", [])
            for mark in marks:
                mt = mark.get("type", "")
                if mt == "bold":
                    text = f"<strong>{text}</strong>"
                elif mt == "italic":
                    text = f"<em>{text}</em>"
                elif mt == "underline":
                    text = f"<u>{text}</u>"
                elif mt == "strike":
                    text = f"<s>{text}</s>"
            if force_italic and not any(m.get("type") == "italic" for m in marks):
                text = f"<em>{text}</em>"
            parts.append(text)
        elif item.get("type") == "hardBreak":
            parts.append("<br>")
    return "".join(parts)


def _tiptap_body_to_html(content: list[dict[str, Any]]) -> str:
    """Convert TipTap JSON body nodes to styled HTML, matching DOCX formatting."""
    html_parts: list[str] = []
    for node in content:
        nt = node.get("type", "")

        if nt == "paragraph":
            inline = _tiptap_inline_to_html(node.get("content", []))
            html_parts.append(
                f'<p style="text-indent:1.25cm; margin-bottom:6px; text-align:justify;">{inline}</p>'
            )

        elif nt == "heading":
            level = node.get("attrs", {}).get("level", 2)
            inline = _tiptap_inline_to_html(node.get("content", []))
            if level == 2:
                html_parts.append(
                    f'<h2 style="font-size:12pt; font-weight:bold; text-transform:uppercase; '
                    f'letter-spacing:0.5px; margin-top:20px; margin-bottom:8px;">{inline}</h2>'
                )
            elif level == 3:
                html_parts.append(
                    f'<h3 style="font-size:12pt; font-weight:bold; margin-top:14px; '
                    f'margin-bottom:6px;">{inline}</h3>'
                )
            else:
                html_parts.append(f'<p>{inline}</p>')

        elif nt == "blockquote":
            bq_parts = []
            for child in node.get("content", []):
                inline = _tiptap_inline_to_html(child.get("content", []), force_italic=True)
                bq_parts.append(f"<p>{inline}</p>")
            html_parts.append(
                f'<blockquote style="margin:10px 2cm 10px 1.2cm; padding:8px 12px; '
                f'border-left:3px solid #666; background:#f5f5f5; font-size:11pt; '
                f'text-align:justify; font-style:italic;">{"".join(bq_parts)}</blockquote>'
            )

        elif nt in ("bulletList", "orderedList"):
            tag = "ul" if nt == "bulletList" else "ol"
            items = []
            for item in node.get("content", []):
                if item.get("type") == "listItem":
                    for child in item.get("content", []):
                        inline = _tiptap_inline_to_html(child.get("content", []))
                        items.append(f"<li>{inline}</li>")
            html_parts.append(f'<{tag} style="margin-left:2em; margin-bottom:10px;">{"".join(items)}</{tag}>')

        else:
            children = node.get("content", [])
            if children:
                html_parts.append(_tiptap_body_to_html(children))

    return "\n".join(html_parts)


def _build_pdf_html(req: ParecerRequest, version: ParecerVersion) -> str:
    """Build complete styled HTML from TipTap JSON, matching DOCX structure exactly."""
    meta = _get_metadata(req)

    # Body content from TipTap JSON (filter out decoration: header, sigs, footer)
    body_html = ""
    tiptap = version.content_tiptap
    if tiptap and isinstance(tiptap, dict):
        body_html = _tiptap_body_to_html(_filter_tiptap_content(tiptap.get("content", [])))
    elif version.content_html:
        body_html = version.content_html

    # Assinaturas
    sigs = ""
    for nome, oab in ADVOGADOS:
        sigs += f"""
        <div class="assinatura">
            <div class="linha-assinatura">
                <div class="nome">{_escape_html(nome)}</div>
                <div class="oab">{oab}</div>
            </div>
        </div>"""

    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<style>
@page {{ size: A4; margin: 2.5cm 2cm 2cm 2cm; }}
* {{ box-sizing: border-box; margin: 0; padding: 0; }}
body {{ font-family: 'Times New Roman', 'Noto Serif', serif; font-size: 12pt;
        line-height: 1.6; color: #000; background: none; }}

.cabecalho {{ border-bottom: 1.5pt solid #000; padding-bottom: 0.35cm; margin-bottom: 0.7cm; }}
.cabecalho .linha1 {{ text-align: center; font-variant: small-caps; font-size: 13pt; letter-spacing: 2px; }}
.cabecalho .linha2 {{ text-align: center; font-weight: bold; font-variant: small-caps; font-size: 11pt;
                      letter-spacing: 1px; padding-top: 3px; }}

h1.titulo {{ text-align: center; font-size: 13pt; font-weight: bold; margin: 14px 0 16px;
             letter-spacing: 1px; }}

.identificacao {{ margin-bottom: 16px; }}
.identificacao p {{ margin: 3px 0; font-size: 11.5pt; }}

.conteudo {{ text-align: justify; }}
.conteudo p {{ font-size: 12pt; margin-bottom: 6px; }}

blockquote {{ break-inside: avoid; }}

.local-data {{ margin-top: 30px; text-align: right; font-size: 12pt; }}

.assinaturas {{ display: flex; flex-wrap: wrap; justify-content: space-between;
                margin-top: 50px; break-inside: avoid; }}
.assinatura {{ width: 46%; text-align: center; margin-bottom: 38px; }}
.assinatura .linha-assinatura {{ border-top: 1px solid #000; margin-top: 42px; padding-top: 5px; }}
.assinatura .nome {{ font-weight: bold; font-variant: small-caps; font-size: 11pt; }}
.assinatura .oab {{ font-size: 10pt; margin-top: 2px; }}

.rodape {{ border-top: 1pt solid #000; padding-top: 0.25cm; margin-top: 1cm;
           font-size: 9pt; text-align: center; color: #333; line-height: 1.4; }}
</style>
</head>
<body>

<!-- CABEÇALHO -->
<div class="cabecalho">
    <div class="linha1">{_escape_html(ESCRITORIO_LINHA1)}</div>
    <div class="linha2">{_escape_html(ESCRITORIO_LINHA2)}</div>
</div>

<!-- TÍTULO -->
<h1 class="titulo">PARECER JURÍDICO</h1>

<!-- IDENTIFICAÇÃO -->
<div class="identificacao">
    <p><strong>Órgão Consulente:</strong> {_escape_html(meta["orgao_consulente"])}</p>
    <p><strong>Referência:</strong> {_escape_html(meta["referencia"])}</p>
    <p><strong>Assunto:</strong> {_escape_html(meta["assunto"])}</p>
</div>

<!-- CORPO -->
<div class="conteudo">
{body_html}
</div>

<!-- LOCAL E DATA -->
<div class="local-data">
    {_escape_html(_formatar_data_extenso(meta["municipio_uf"]))}
</div>

<!-- ASSINATURAS -->
<div class="assinaturas">
{sigs}
</div>

<!-- RODAPÉ -->
<div class="rodape">
    {_escape_html(ESCRITORIO_ENDERECO)}
</div>

</body>
</html>"""


async def to_pdf(
    parecer_request: ParecerRequest,
    version: ParecerVersion,
    db: AsyncSession,
) -> bytes:
    """Generate PDF from TipTap JSON — same structure as DOCX."""
    from weasyprint import HTML

    html_str = _build_pdf_html(parecer_request, version)
    return HTML(string=html_str).write_pdf()
